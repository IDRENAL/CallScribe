"""Экспорт стенограммы (+ выжимки) в .docx (v2.0).

Стенограмма строится из структурированного JSON (надёжнее, чем парсить Markdown):
шапка с метаданными, выжимка (если есть) и реплики с таймкодами и спикерами.
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from transcriber import format_timestamp

ACCENT = RGBColor(0x00, 0x99, 0x66)   # таймкоды/акценты
MUTED = RGBColor(0x66, 0x66, 0x66)


def _render_inline(paragraph, text: str) -> None:
    """Минимальный inline-Markdown: **жирный** → bold-run, остальное обычным."""
    for i, part in enumerate(text.split("**")):
        if not part:
            continue
        run = paragraph.add_run(part)
        if i % 2 == 1:           # нечётные куски — внутри **…**
            run.bold = True


def _add_summary(doc: Document, summary_md: str) -> None:
    """Вставить выжимку, разбирая её упрощённый Markdown."""
    doc.add_heading("Выжимка", level=1)
    for raw in summary_md.splitlines():
        line = raw.rstrip()
        if not line or line == "---":
            continue
        if line.startswith("# "):           # заголовок самой выжимки — пропускаем
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _render_inline(p, line[2:].strip())
        else:
            p = doc.add_paragraph()
            _render_inline(p, line)


def _add_meta(doc: Document, data: dict) -> None:
    src = Path(data.get("source_file", "")).name
    date = (data.get("transcribed_at", "") or "").replace("T", " ")[:16]
    dur = (data.get("duration_seconds") or 0) / 60
    lang = data.get("language", "")
    prob = data.get("language_probability")
    lang_str = f"{lang} ({prob * 100:.0f}%)" if prob else lang
    rows = [("Файл", src), ("Дата", date), ("Длительность", f"{dur:.1f} мин"),
            ("Язык", lang_str), ("Модель", f"whisper-{data.get('model', '')}")]
    for label, value in rows:
        if not value:
            continue
        p = doc.add_paragraph()
        r = p.add_run(f"{label}: ")
        r.bold = True
        p.add_run(str(value))


def _add_transcript(doc: Document, data: dict) -> None:
    doc.add_heading("Стенограмма", level=1)
    for seg in data.get("segments", []):
        p = doc.add_paragraph()
        ts = p.add_run(f"[{format_timestamp(seg['start'])}] ")
        ts.bold = True
        ts.font.color.rgb = ACCENT
        spk = seg.get("speaker")
        if spk:
            sr = p.add_run(f"{spk}: ")
            sr.bold = True
        p.add_run(seg.get("text", ""))


def export_to_docx(data: dict, summary_md: str | None, out_path: str | Path) -> Path:
    """Собрать .docx из данных стенограммы и (опц.) текста выжимки."""
    doc = Document()
    title = doc.add_heading("Стенограмма звонка", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    _add_meta(doc, data)
    if summary_md and summary_md.strip():
        _add_summary(doc, summary_md)
    _add_transcript(doc, data)

    out_path = Path(out_path)
    doc.save(str(out_path))
    return out_path


def build_export(stem: str, transcripts_dir: str | Path) -> Path:
    """Найти стенограмму/выжимку по stem и выгрузить <stem>.docx рядом."""
    transcripts_dir = Path(transcripts_dir)
    json_path = transcripts_dir / f"{stem}_transcript.json"
    if not json_path.exists():
        raise FileNotFoundError(f"Нет стенограммы для экспорта: {stem}")
    data = json.loads(json_path.read_text(encoding="utf-8"))
    summary_path = transcripts_dir / f"{stem}_summary.md"
    summary_md = summary_path.read_text(encoding="utf-8") if summary_path.exists() else None
    out = transcripts_dir / f"{stem}.docx"
    return export_to_docx(data, summary_md, out)
