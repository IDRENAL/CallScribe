"""Тесты exporter.py — сборка .docx из данных стенограммы и выжимки."""
from __future__ import annotations

import json

import pytest

from docx import Document

from exporter import build_export, export_to_docx

DATA = {
    "source_file": "recordings/call_2026-05-30_09-20-00.wav",
    "transcribed_at": "2026-05-30T09:25:00",
    "language": "ru", "language_probability": 0.98,
    "duration_seconds": 125.0, "model": "large-v3", "mode": "accurate",
    "segments": [
        {"start": 0.0, "end": 2.0, "text": "Привет", "speaker": "Я"},
        {"start": 2.0, "end": 5.0, "text": "Здравствуйте", "speaker": "Собеседник"},
        {"start": 65.0, "end": 67.0, "text": "Без спикера"},
    ],
}


def _all_text(path) -> str:
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def test_export_creates_valid_docx(tmp_path):
    out = export_to_docx(DATA, None, tmp_path / "call.docx")
    assert out.exists()
    text = _all_text(out)
    assert "Стенограмма звонка" in text   # заголовок
    assert "Привет" in text
    assert "Здравствуйте" in text


def test_export_includes_timestamps_and_speakers(tmp_path):
    out = export_to_docx(DATA, None, tmp_path / "c.docx")
    text = _all_text(out)
    assert "[00:00]" in text
    assert "[01:05]" in text              # 65 сек → 01:05
    assert "Я:" in text and "Собеседник:" in text


def test_export_includes_meta(tmp_path):
    out = export_to_docx(DATA, None, tmp_path / "c.docx")
    text = _all_text(out)
    assert "call_2026-05-30_09-20-00.wav" in text   # только имя файла
    assert "whisper-large-v3" in text
    assert "2.1 мин" in text                          # 125 сек


def test_export_embeds_summary_when_present(tmp_path):
    summary = ("# Выжимка звонка\n\n## Краткое содержание\n"
               "Обсудили **релиз**\n\n## Задачи\n- Сделать бэкап\n")
    out = export_to_docx(DATA, summary, tmp_path / "c.docx")
    text = _all_text(out)
    assert "Выжимка" in text
    assert "Краткое содержание" in text
    assert "Обсудили релиз" in text       # ** убраны, текст сохранён
    assert "Сделать бэкап" in text        # пункт списка


def test_build_export_reads_files_by_stem(tmp_path):
    stem = "call_x"
    (tmp_path / f"{stem}_transcript.json").write_text(
        json.dumps(DATA, ensure_ascii=False), encoding="utf-8")
    (tmp_path / f"{stem}_summary.md").write_text(
        "## Краткое содержание\nитог встречи", encoding="utf-8")
    out = build_export(stem, tmp_path)
    assert out == tmp_path / f"{stem}.docx"
    text = _all_text(out)
    assert "итог встречи" in text
    assert "Привет" in text


def test_build_export_missing_transcript_raises(tmp_path):
    with pytest.raises(FileNotFoundError):
        build_export("nope", tmp_path)


def test_export_without_summary_has_no_summary_heading(tmp_path):
    out = export_to_docx(DATA, None, tmp_path / "c.docx")
    doc = Document(str(out))
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "Выжимка" not in headings
    assert "Стенограмма" in headings
